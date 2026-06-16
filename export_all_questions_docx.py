from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


SCRIPT = Path("script.js")
OUTPUT = Path("جميع_أسئلة_مسابقة_الحروف_مرتبة.docx")

ARABIC_LETTER_ORDER = [
    "أ", "ب", "ت", "ث", "ج", "ح", "خ", "د", "ذ", "ر", "ز", "س", "ش", "ص",
    "ض", "ط", "ظ", "ع", "غ", "ف", "ق", "ك", "ل", "م", "ن", "هـ", "و", "ي",
]

DIFFICULTY_ORDER = {
    "simple": 0,
    "medium": 1,
    "aboveMedium": 2,
}


def extract_book_letters() -> list[dict]:
    script = SCRIPT.read_text(encoding="utf-8")
    start = script.index("const BOOK_LETTERS = ") + len("const BOOK_LETTERS = ")
    end = script.index(";\n\nfunction buildBookQuestionBank", start)
    return json.loads(script[start:end])


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)


def add_summary(document: Document, letters: list[dict]) -> None:
    total = sum(len(item["questions"]) for item in letters)
    title = document.add_heading("جميع أسئلة مسابقة الحروف", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    summary = document.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = summary.add_run(f"إجمالي الأسئلة: {total} سؤال | عدد الحروف: {len(letters)}")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12)


def add_letter_table(document: Document, letter: str, questions: list[dict]) -> None:
    heading = document.add_heading(f"حرف {letter} ({len(questions)} سؤال)", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    sorted_questions = sorted(
        enumerate(questions, start=1),
        key=lambda pair: (
            DIFFICULTY_ORDER.get(pair[1].get("difficulty"), 99),
            pair[0],
        ),
    )

    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    headers = ["م", "المستوى", "السؤال", "الإجابة"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True)

    for number, question in sorted_questions:
        row = table.add_row().cells
        set_cell_text(row[0], number)
        set_cell_text(row[1], question.get("difficultyLabel", ""))
        set_cell_text(row[2], question.get("question", ""))
        set_cell_text(row[3], question.get("answer", ""))

    document.add_paragraph()


def main() -> None:
    letters = extract_book_letters()
    by_letter = {item["letter"]: item for item in letters}
    ordered_letters = [by_letter[letter] for letter in ARABIC_LETTER_ORDER if letter in by_letter]

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    add_summary(document, ordered_letters)

    for index, item in enumerate(ordered_letters):
      add_letter_table(document, item["letter"], item["questions"])
      if index < len(ordered_letters) - 1:
          document.add_page_break()

    document.save(OUTPUT)
    print(f"output={OUTPUT.resolve()}")
    print(f"letters={len(ordered_letters)}")
    print(f"questions={sum(len(item['questions']) for item in ordered_letters)}")


if __name__ == "__main__":
    main()
