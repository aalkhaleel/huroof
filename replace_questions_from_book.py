from __future__ import annotations

import json
import re
from collections import Counter
from pathlib import Path

from docx import Document


SCRIPT = Path("script.js")
REPORT = Path("book_questions_import_report.txt")

SOURCE_FILES = [
    Path.home() / "Desktop" / "كتاب مسابقات حروف.docx",
    Path(__file__).with_name("أسئلة_حروف.docx"),
]

ARABIC_LETTERS = [
    "أ", "ب", "ت", "ث", "ج", "ح", "خ", "د", "ذ", "ر", "ز", "س", "ش", "ص",
    "ض", "ط", "ظ", "ع", "غ", "ف", "ق", "ك", "ل", "م", "ن", "هـ", "و", "ي",
]

APP_LETTER_ORDER = [
    "ن", "ج", "ل", "خ", "ط", "ث", "ع", "ظ", "ر", "ص", "م", "ض", "س", "ش",
    "ذ", "د", "أ", "ب", "ف", "غ", "ك", "ت", "ح", "ز", "ق", "هـ", "و", "ي",
]

LEVELS = [
    ("simple", "بسيط", "المستوى الأول"),
    ("medium", "متوسط", "المستوى الثاني"),
    ("aboveMedium", "فوق المتوسط", "المستوى الثالث"),
]


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\xa0", " ")).strip()


def normalize_letter(letter: str) -> str:
    letter = clean(letter)
    return {
        "ه": "هـ",
        "ة": "هـ",
        "ا": "أ",
        "إ": "أ",
        "آ": "أ",
    }.get(letter, letter)


def question_key(question: str, answer: str) -> str:
    return f"{clean(question).casefold()}::{clean(answer).casefold()}"


def table_cells(row) -> list[str]:
    return [clean(cell.text) for cell in row.cells]


def detect_level(index: int, total: int) -> tuple[str, str, str]:
    if total <= 0:
        return LEVELS[1]

    ratio = index / total
    if ratio < 1 / 3:
        return LEVELS[0]
    if ratio < 2 / 3:
        return LEVELS[1]
    return LEVELS[2]


def append_question(
    by_letter: dict[str, list[dict]],
    seen: set[str],
    letter: str,
    question: str,
    answer: str,
    level_key: str,
    level_label: str,
    level_subject: str,
) -> bool:
    letter = normalize_letter(letter)
    if letter not in by_letter:
        return False

    key = question_key(question, answer)
    if key in seen:
        return False

    seen.add(key)
    by_letter[letter].append({
        "subject": level_subject,
        "question": clean(question),
        "answer": clean(answer),
        "difficulty": level_key,
        "difficultyLabel": level_label,
    })
    return True


def extract_from_tables(document: Document, by_letter: dict[str, list[dict]], seen: set[str]) -> Counter:
    counts: Counter = Counter()

    for table_index, table in enumerate(document.tables):
        if table_index // len(LEVELS) >= len(ARABIC_LETTERS):
            continue

        letter = ARABIC_LETTERS[table_index // len(LEVELS)]
        level_key, level_label, level_subject = LEVELS[table_index % len(LEVELS)]

        for row in table.rows:
            cells = table_cells(row)
            if len(cells) < 3:
                continue

            number, question, answer = cells[0], cells[1], cells[2]
            if not question or not answer or not re.search(r"\d", number):
                continue

            if append_question(by_letter, seen, letter, question, answer, level_key, level_label, level_subject):
                counts[letter] += 1

    return counts


def extract_from_paragraphs(document: Document, by_letter: dict[str, list[dict]], seen: set[str]) -> Counter:
    raw_by_letter: dict[str, list[tuple[str, str]]] = {letter: [] for letter in ARABIC_LETTERS}
    current_letter: str | None = None
    pending_question: str | None = None

    for paragraph in document.paragraphs:
        text = clean(paragraph.text)
        if not text:
            continue

        if text.startswith("حرف") and "(" in text:
            heading = text.replace("حرف", "", 1).split("(", 1)[0].strip()
            current_letter = normalize_letter(heading)
            pending_question = None
            continue

        if current_letter not in raw_by_letter:
            continue

        numbered = re.match(r"^\d+\.\s*(.+)$", text)
        if numbered:
            pending_question = clean(numbered.group(1))
            continue

        if text.startswith("الإجابة:") and pending_question:
            answer = clean(text.split(":", 1)[1])
            raw_by_letter[current_letter].append((pending_question, answer))
            pending_question = None

    counts: Counter = Counter()
    for letter, rows in raw_by_letter.items():
        total = len(rows)
        for index, (question, answer) in enumerate(rows):
            level_key, level_label, level_subject = detect_level(index, total)
            if append_question(by_letter, seen, letter, question, answer, level_key, level_label, level_subject):
                counts[letter] += 1

    return counts


def extract_questions() -> tuple[list[dict], Counter, dict[str, Counter]]:
    by_letter: dict[str, list[dict]] = {letter: [] for letter in ARABIC_LETTERS}
    seen: set[str] = set()
    source_counts: dict[str, Counter] = {}

    for source in SOURCE_FILES:
        if not source.exists():
            source_counts[str(source)] = Counter()
            continue

        document = Document(source)
        if document.tables:
            counts = extract_from_tables(document, by_letter, seen)
        else:
            counts = extract_from_paragraphs(document, by_letter, seen)
        source_counts[str(source)] = counts

    app_letters = []
    app_counts: Counter = Counter()
    for letter in APP_LETTER_ORDER:
        questions = by_letter.get(letter, [])
        if not questions:
            continue
        app_letters.append({"letter": letter, "questions": questions})
        app_counts[letter] = len(questions)

    return app_letters, app_counts, source_counts


def make_js(app_letters: list[dict]) -> str:
    data = json.dumps(app_letters, ensure_ascii=False, indent=2)
    return f"""const BOOK_LETTERS = {data};

function buildBookQuestionBank(bookLetters) {{
  return bookLetters.map((item) => ({{
    ...item,
    questions: item.questions.map((question) => ({{
      ...question,
      subject: question.subject || "كتاب مسابقات حروف",
      difficulty: question.difficulty || "medium",
      difficultyLabel: question.difficultyLabel || "متوسط",
    }})),
  }}));
}}

const LETTERS = buildBookQuestionBank(BOOK_LETTERS);
"""


def replace_question_block(new_block: str) -> None:
    script = SCRIPT.read_text(encoding="utf-8")
    start = script.find("const BASE_LETTERS = [")
    if start == -1:
        start = script.index("const BOOK_LETTERS = [")
    end = script.index("const state = {")
    SCRIPT.write_text(script[:start] + new_block + "\n" + script[end:], encoding="utf-8")


def write_report(app_counts: Counter, source_counts: dict[str, Counter]) -> None:
    lines = [
        "sources:",
        *[f"{source}: {sum(counts.values())}" for source, counts in source_counts.items()],
        f"total_used_in_app={sum(app_counts.values())}",
        "used_counts:",
        *[f"{letter}: {app_counts[letter]}" for letter in APP_LETTER_ORDER if app_counts[letter]],
        "omitted_because_board_rotates_25_of_28_letters:",
    ]
    REPORT.write_text("\n".join(lines), encoding="utf-8")


if __name__ == "__main__":
    letters, app_counts, source_counts = extract_questions()
    replace_question_block(make_js(letters))
    write_report(app_counts, source_counts)
    print(f"source_total={sum(sum(counts.values()) for counts in source_counts.values())}")
    print(f"total_used_in_app={sum(app_counts.values())}")
    print(f"letters_used={len(letters)}")
