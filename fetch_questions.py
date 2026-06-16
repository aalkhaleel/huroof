import sys, requests, time
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding="utf-8")

BASE = "https://sahamstudio.com/horoof"
SUBSCRIPTION = "4070530617"
PASSWORD = "1234"
OUTPUT_FILE = r"c:\huroof\أسئلة_حروف.docx"

LETTERS = ['أ','ب','ت','ث','ج','ح','خ','د','ذ','ر','ز','س','ش','ص','ض','ط','ظ','ع','غ','ف','ق','ك','ل','م','ن','هـ','و','ي']

# ── تسجيل الدخول ────────────────────────────────────────────────
session = requests.Session()
session.headers.update({
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
    "Referer": BASE + "/questions.php",
})

print("جاري تسجيل الدخول...")
session.post(BASE + "/verify.php", data={"code": SUBSCRIPTION, "password": PASSWORD})
print("تم تسجيل الدخول.\n")

# ── جلب الأسئلة ─────────────────────────────────────────────────
all_questions = {}  # letter -> list of {id, question, answer}

for letter in LETTERS:
    seen_ids = set()
    questions = []
    reset = True

    print(f"حرف [{letter}] ...", end=" ", flush=True)

    while True:
        params = {"letter": letter}
        if reset:
            params["reset"] = "1"
            reset = False

        try:
            r = session.get(BASE + "/get_question.php", params=params, timeout=10)
            data = r.json()
        except Exception as e:
            print(f"خطأ: {e}")
            break

        if data.get("status") != "success":
            break

        qid = data["id"]

        # توقف عند التكرار
        if qid in seen_ids:
            break

        seen_ids.add(qid)
        questions.append({
            "id": qid,
            "question": data["question"],
            "answer": data["answer"],
        })

        # حد أقصى 200 سؤال لكل حرف لتفادي الحلقة اللانهائية
        if len(questions) >= 200:
            break

        time.sleep(0.15)

    all_questions[letter] = questions
    print(f"{len(questions)} سؤال")

total = sum(len(v) for v in all_questions.values())
print(f"\nإجمالي الأسئلة: {total}")

# ── إنشاء ملف Word ──────────────────────────────────────────────
print("\nجاري إنشاء ملف Word...")
doc = Document()

# ضبط الهوامش
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# عنوان رئيسي
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("أسئلة حروف مع عزيز")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x10, 0x1a, 0x13)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run(f"إجمالي الأسئلة: {total} سؤال | رقم الاشتراك: {SUBSCRIPTION}").font.size = Pt(11)

doc.add_paragraph()

# الأسئلة مرتبة حسب الحرف
for letter in LETTERS:
    questions = all_questions.get(letter, [])
    if not questions:
        continue

    # عنوان الحرف
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h_run = heading.add_run(f"حرف  {letter}  ({len(questions)} سؤال)")
    h_run.bold = True
    h_run.font.size = Pt(16)
    h_run.font.color.rgb = RGBColor(0x00, 0x8f, 0x00)

    # فاصل
    hr = doc.add_paragraph()
    hr_run = hr.add_run("─" * 60)
    hr_run.font.size = Pt(9)
    hr_run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)

    for i, q in enumerate(questions, 1):
        # رقم + سؤال
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(2)
        num = p.add_run(f"{i}. ")
        num.bold = True
        num.font.size = Pt(11)
        num.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        q_run = p.add_run(q["question"])
        q_run.font.size = Pt(11)

        # الإجابة
        a = doc.add_paragraph()
        a.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        a.paragraph_format.space_after = Pt(8)
        a.paragraph_format.left_indent = Cm(1)
        a_label = a.add_run("الإجابة: ")
        a_label.bold = True
        a_label.font.size = Pt(11)
        a_label.font.color.rgb = RGBColor(0x00, 0x6a, 0xc1)
        a_run = a.add_run(q["answer"])
        a_run.font.size = Pt(11)
        a_run.bold = True
        a_run.font.color.rgb = RGBColor(0x00, 0x6a, 0xc1)

    doc.add_paragraph()

doc.save(OUTPUT_FILE)
print(f"تم حفظ الملف: {OUTPUT_FILE}")
