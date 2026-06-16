from __future__ import annotations

import html
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


SOURCE = Path("video_transcript.ar-orig.vtt")
CLEAN_TEXT = Path("نص_الفيديو_منظف.txt")
OUTPUT = Path("أسئلة_مستخرجة_من_الفيديو.docx")
OUTPUT_FIXED = Path("أسئلة_مستخرجة_من_الفيديو_متوافق.docx")

LETTER_NAMES = {
    "الالف": "أ", "الألف": "أ", "الف": "أ",
    "الباء": "ب", "باء": "ب",
    "التاء": "ت", "تاء": "ت",
    "الثاء": "ث", "ثاء": "ث",
    "الجيم": "ج", "جيم": "ج",
    "الحاء": "ح", "حاء": "ح",
    "الخاء": "خ", "خاء": "خ",
    "الدال": "د", "دال": "د",
    "الذال": "ذ", "ذال": "ذ",
    "الراء": "ر", "راء": "ر",
    "الزاي": "ز", "زاي": "ز",
    "السين": "س", "سين": "س",
    "الشين": "ش", "شين": "ش",
    "الصاد": "ص", "صاد": "ص",
    "الضاد": "ض", "ضاد": "ض",
    "الطاء": "ط", "طاء": "ط",
    "الظاء": "ظ", "ظاء": "ظ",
    "العين": "ع", "عين": "ع",
    "الغين": "غ", "غين": "غ",
    "الفاء": "ف", "فاء": "ف",
    "القاف": "ق", "قاف": "ق",
    "الكاف": "ك", "كاف": "ك",
    "اللام": "ل", "لام": "ل",
    "الميم": "م", "ميم": "م",
    "النون": "ن", "نون": "ن",
    "الهاء": "هـ", "هاء": "هـ",
    "الواو": "و", "واو": "و",
    "الياء": "ي", "ياء": "ي",
}

STOP_MARKERS = [
    "احسنت", "أحسنت", "انتهى الوقت", "انتهى", "الفريق الاخضر",
    "الفريق الأحمر", "الفريق الاحمر", "اختيار", "وننتقل", "وما نزال",
]


def clean_vtt(path: Path) -> str:
    text = path.read_text(encoding="utf-8")
    text = re.sub(r"<[^>]+>", "", text)
    lines: list[str] = []
    previous = ""
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        if "-->" in line:
            continue
        if line.startswith(("WEBVTT", "Kind:", "Language:")):
            continue
        if line.startswith("[") and line.endswith("]"):
            continue
        line = re.sub(r"\s+", " ", line)
        if line == previous:
            continue
        lines.append(line)
        previous = line
    clean = re.sub(r"\s+", " ", " ".join(lines)).strip()
    CLEAN_TEXT.write_text(clean, encoding="utf-8")
    return clean


def normalize_space(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip(" ،.:-")


def find_letter(segment: str) -> str:
    matches = list(re.finditer(r"حرف\s+([اأإآ]?\w+)", segment))
    if not matches:
        return ""
    raw = matches[-1].group(1).replace("اا", "ا")
    return LETTER_NAMES.get(raw, LETTER_NAMES.get(raw.replace("أ", "ا"), raw[:1]))


def find_category(segment: str) -> str:
    match = re.search(r"في مجال\s+([^،.]{3,35})", segment)
    if not match:
        return ""
    category = normalize_space(match.group(1))
    for stop in ["والسؤال", "السؤال", "وهو", "هو"]:
        if stop in category:
            category = category.split(stop)[0]
    return normalize_space(category)


def trim_question(segment: str) -> str:
    starters = [
        "السؤال يقول", "والسؤال يقول", "السؤال التالي", "هذا السؤال",
        "والسؤال هو التالي", "وسؤال يقول", "وسؤال هو التالي",
    ]
    start = 0
    for marker in starters:
        idx = segment.rfind(marker)
        if idx != -1:
            start = idx + len(marker)
            break
    question = segment[start:]
    question = re.sub(r"^.*?في مجال\s+[^،.]{2,35}", "", question)
    question = re.split(r"\b(?:ايوه|نعم)\b", question, maxsplit=1)[0]
    return normalize_space(question)


def find_answer(segment: str) -> str:
    match = re.search(r"(?:ايوه|نعم)\s+(.{2,90})", segment)
    if not match:
        return ""
    answer = match.group(1)
    answer = answer.split("ايوه")[0].split("نعم")[0]
    for marker in STOP_MARKERS:
        if marker in answer:
            answer = answer.split(marker)[0]
    answer = re.sub(r"\b(?:يا|لا)\b.*$", "", answer).strip()
    words = answer.split()
    if len(words) > 5:
        answer = " ".join(words[:5])
    return normalize_space(answer)


def extract_questions(text: str) -> list[dict[str, str]]:
    marker_re = r"(?:والسؤال\s+يقول|السؤال\s+يقول|السؤال\s+التالي|هذا\s+السؤال|وهذا\s+السؤال|والسؤال\s+هو\s+التالي|وسؤال\s+يقول|وسؤال\s+هو\s+التالي)"
    positions = [m.start() for m in re.finditer(marker_re, text)]
    items: list[dict[str, str]] = []
    seen_questions: set[str] = set()

    for idx, pos in enumerate(positions):
        end = positions[idx + 1] if idx + 1 < len(positions) else min(len(text), pos + 900)
        context_start = max(0, pos - 180)
        context = normalize_space(text[context_start:end])
        segment = normalize_space(text[pos:end])
        question = trim_question(segment)
        answer = find_answer(segment)
        if len(question) < 12 or len(question) > 260:
            continue
        if any(bad in question for bad in [
            "يقرع الجرس",
            "الاجابه الاولى",
            "المسابقة",
            "يستطيع الاجابه",
            "ساطرح",
            "سأطرح",
            "ملاحظه",
            "ملاحظة",
            "الحق في الاجابه",
        ]):
            continue
        key = question[:90]
        if key in seen_questions:
            continue
        seen_questions.add(key)
        items.append({
            "letter": find_letter(context),
            "category": find_category(context),
            "question": question,
            "answer": answer,
            "source": context[:420],
        })

    return items


def esc(text: object) -> str:
    return html.escape(str(text), quote=True)


def p(text: str, style: str | None = None) -> str:
    style_xml = f'<w:pPr><w:pStyle w:val="{style}"/></w:pPr>' if style else ""
    return f'<w:p>{style_xml}<w:r><w:rPr><w:rtl/></w:rPr><w:t>{esc(text)}</w:t></w:r></w:p>'


def tc(text: str) -> str:
    return f'<w:tc><w:tcPr><w:textDirection w:val="rlTb"/></w:tcPr>{p(text)}</w:tc>'


def tr(values: list[str]) -> str:
    return "<w:tr>" + "".join(tc(value) for value in values) + "</w:tr>"


def write_docx(items: list[dict[str, str]], path: Path) -> None:
    rows = [tr(["م", "الحرف", "المجال", "السؤال", "الإجابة من الفيديو", "مقطع للمراجعة"])]
    for index, item in enumerate(items, start=1):
        rows.append(tr([
            str(index),
            item["letter"],
            item["category"],
            item["question"],
            item["answer"],
            item["source"],
        ]))

    document = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body>'
        + p("أسئلة مستخرجة من فيديو يوتيوب", "Title")
        + p("تم استخراج هذه الأسئلة آليا من الترجمة العربية، لذلك أضفت عمود مقطع للمراجعة لتدقيق النص والإجابة.")
        + "<w:tbl>"
        + "".join(rows)
        + "</w:tbl>"
        + '<w:sectPr><w:bidi/><w:pgSz w:w="16838" w:h="11906" w:orient="landscape"/>'
        + '<w:pgMar w:top="720" w:right="720" w:bottom="720" w:left="720"/></w:sectPr>'
        + "</w:body></w:document>"
    )
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
</Types>"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""
    styles = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:styleId="Normal"><w:name w:val="Normal"/><w:pPr><w:bidi/></w:pPr><w:rPr><w:rtl/><w:sz w:val="20"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="Title"><w:name w:val="Title"/><w:pPr><w:bidi/><w:jc w:val="center"/></w:pPr><w:rPr><w:rtl/><w:b/><w:sz w:val="34"/></w:rPr></w:style>
</w:styles>"""
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as docx:
        docx.writestr("[Content_Types].xml", content_types)
        docx.writestr("_rels/.rels", rels)
        docx.writestr("word/document.xml", document)
        docx.writestr("word/styles.xml", styles)


def write_docx_compatible(items: list[dict[str, str]], path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    title = document.add_heading("أسئلة مستخرجة من فيديو يوتيوب", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    note = document.add_paragraph(
        "تم استخراج هذه الأسئلة آليا من الترجمة العربية للفيديو. "
        "يرجى مراجعة عمود المقطع قبل اعتماد السؤال، لأن الترجمة التلقائية قد تحتوي على تكرار أو أخطاء."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    headers = ["م", "الحرف", "المجال", "السؤال", "الإجابة من الفيديو", "مقطع للمراجعة"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header

    for index, item in enumerate(items, start=1):
        row = table.add_row().cells
        values = [
            str(index),
            item["letter"],
            item["category"],
            item["question"],
            item["answer"],
            item["source"],
        ]
        for cell, value in zip(row, values):
            cell.text = value

    document.save(path)


if __name__ == "__main__":
    clean = clean_vtt(SOURCE)
    questions = extract_questions(clean)
    write_docx_compatible(questions, OUTPUT_FIXED)
    print(f"questions={len(questions)}")
    print(str(OUTPUT_FIXED.resolve()).encode("unicode_escape").decode("ascii"))
