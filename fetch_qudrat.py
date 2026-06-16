import sys, requests, time, re
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding="utf-8")

OUTPUT_FILE = r"c:\huroof\أسئلة_قدرات.docx"

HEADERS = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}

# Pages to scrape — (section title, URL)
PAGES = [
    ("أسئلة قدرات لفظي مع الإجابة",     "https://as2elah.com/1214/"),
    ("أسئلة قدرات رياضيات",             "https://as2elah.com/1216/"),
    ("أسئلة قدرات كمي محلولة",          "https://as2elah.com/1202/"),
    ("أسئلة قدرات كمي pdf مع الإجابات", "https://as2elah.com/1208/"),
    ("أسئلة قدرات ذهنية مع الحل",       "https://as2elah.com/1192/"),
    ("أسئلة قدرات — المرجع",            "https://almrj3.com/ability-questions/"),
    ("أسئلة قدرات — روتيسا",            "https://routesa.app/أسئلة-قدرات-مع-الإجابة/"),
]

session = requests.Session()
session.headers.update(HEADERS)

# ── helpers ──────────────────────────────────────────────────────────────────

def clean(text):
    return re.sub(r"\s+", " ", text or "").strip()

def get_soup(url, retries=3):
    for attempt in range(retries):
        try:
            r = session.get(url, timeout=25)
            r.encoding = "utf-8"
            return BeautifulSoup(r.text, "html.parser")
        except Exception as e:
            if attempt < retries - 1:
                time.sleep(3)
            else:
                print(f"  خطأ: {e}")
    return None

# ── parsers ───────────────────────────────────────────────────────────────────

def parse_table_format(soup):
    """Verbal format: <table> question | options | correct answer"""
    questions = []
    for table in soup.find_all("table"):
        rows = table.find_all("tr")
        for row in rows:
            cells = row.find_all("td")
            if len(cells) < 2:
                continue
            q_text = clean(cells[0].get_text())
            if not q_text or q_text in ("السؤال", "سؤال", "العبارة"):
                continue
            options = [clean(li.get_text()) for li in cells[1].find_all("li")]
            answer = clean(cells[2].get_text()) if len(cells) >= 3 else ""
            if q_text:
                questions.append({"question": q_text, "options": options, "answer": answer})
    return questions

def parse_list_format(soup):
    """<li> with السؤال: ... الإجابة: ..."""
    questions = []
    for li in soup.find_all("li"):
        text = li.get_text(" ", strip=True)
        if "السؤال" not in text and "الإجابة" not in text:
            continue
        q_match = re.search(r"السؤال[:\s]*(.+?)(?:الإجابة|$)", text, re.DOTALL)
        a_match = re.search(r"الإجابة[:\s]*(.+)", text, re.DOTALL)
        q_text = clean(q_match.group(1)) if q_match else ""
        a_text = clean(a_match.group(1)) if a_match else ""
        if q_text and len(q_text) > 5:
            questions.append({"question": q_text, "options": [], "answer": a_text})
    return questions

def parse_numbered_format(soup):
    """Numbered items: 1. question ... أ. choice ... الإجابة: X"""
    questions = []
    article = (soup.find("article") or
               soup.find("div", class_=re.compile(r"entry|content|post")) or
               soup.find("main"))
    if not article:
        return questions
    full_text = article.get_text("\n", strip=True)
    lines = [l.strip() for l in full_text.splitlines() if l.strip()]
    i = 0
    while i < len(lines):
        line = lines[i]
        if re.match(r"^\d+[\.)\-\s]", line) and len(line) > 10:
            q_text = re.sub(r"^\d+[\.)\-]\s*", "", line).strip()
            options = []
            answer = ""
            i += 1
            while i < len(lines):
                nxt = lines[i]
                if re.match(r"^[أبجد][\.)\-]", nxt):
                    options.append(re.sub(r"^[أبجد][\.)\-]\s*", "", nxt))
                    i += 1
                elif re.match(r"^[abcd][\.)\-]", nxt, re.I):
                    options.append(re.sub(r"^[abcd][\.)\-]\s*", "", nxt, flags=re.I))
                    i += 1
                elif re.search(r"(?:الإجابة|الجواب|الصحيح)[:\s]*(.+)", nxt):
                    m = re.search(r"(?:الإجابة|الجواب|الصحيح)[:\s]*(.+)", nxt)
                    answer = clean(m.group(1))
                    i += 1
                    break
                elif re.match(r"^\d+[\.)\-\s]", nxt) and len(nxt) > 10:
                    break
                else:
                    i += 1
            if q_text and len(q_text) > 10:
                questions.append({"question": q_text, "options": options, "answer": answer})
        else:
            i += 1
    return questions

def scrape_page(url):
    soup = get_soup(url)
    if not soup:
        return []
    results = parse_table_format(soup)
    if results:
        return results
    results = parse_list_format(soup)
    if results:
        return results
    return parse_numbered_format(soup)

def get_all_pages(base_url):
    """Follow /page/N/ pagination."""
    all_q = []
    visited = set()
    queue = [base_url]
    while queue:
        url = queue.pop(0)
        if url in visited:
            continue
        visited.add(url)
        print(f"    → {url}")
        qs = scrape_page(url)
        all_q.extend(qs)
        # Find next page links
        soup = get_soup(url)
        if soup:
            for a in soup.find_all("a", href=True):
                href = a["href"]
                if "/page/" in href and href not in visited:
                    queue.append(href)
        time.sleep(0.6)
    return all_q

# ── main ──────────────────────────────────────────────────────────────────────

sections = []
total = 0

for title, url in PAGES:
    print(f"\n[{title}]")
    qs = get_all_pages(url)
    seen = set()
    unique = []
    for q in qs:
        key = q["question"][:70]
        if key not in seen:
            seen.add(key)
            unique.append(q)
    sections.append((title, unique))
    total += len(unique)
    print(f"  {len(unique)} سؤال")

print(f"\nإجمالي الأسئلة: {total}")

# ── Word document ─────────────────────────────────────────────────────────────
print("\nجاري إنشاء ملف Word...")
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("أسئلة اختبار القدرات")
r.bold = True; r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x10, 0x1a, 0x13)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run(f"إجمالي الأسئلة: {total} سؤال").font.size = Pt(11)

doc.add_paragraph()

for sec_title, questions in sections:
    if not questions:
        continue

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h = heading.add_run(f"{sec_title}  ({len(questions)} سؤال)")
    h.bold = True; h.font.size = Pt(15)
    h.font.color.rgb = RGBColor(0x00, 0x6a, 0x00)

    hr = doc.add_paragraph()
    hr.add_run("─" * 60).font.size = Pt(9)

    for i, q in enumerate(questions, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(2)
        num = p.add_run(f"{i}. ")
        num.bold = True; num.font.size = Pt(11)
        num.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        qr = p.add_run(q["question"])
        qr.font.size = Pt(11)

        for opt in q["options"]:
            op = doc.add_paragraph()
            op.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            op.paragraph_format.space_after = Pt(1)
            op.paragraph_format.left_indent = Cm(1)
            op.add_run(f"• {opt}").font.size = Pt(10.5)

        if q["answer"]:
            a = doc.add_paragraph()
            a.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            a.paragraph_format.space_after = Pt(8)
            a.paragraph_format.left_indent = Cm(1)
            al = a.add_run("الإجابة: ")
            al.bold = True; al.font.size = Pt(11)
            al.font.color.rgb = RGBColor(0x00, 0x6a, 0xc1)
            av = a.add_run(q["answer"])
            av.bold = True; av.font.size = Pt(11)
            av.font.color.rgb = RGBColor(0x00, 0x6a, 0xc1)
        else:
            doc.add_paragraph().paragraph_format.space_after = Pt(8)

    doc.add_paragraph()

doc.save(OUTPUT_FILE)
print(f"تم حفظ الملف: {OUTPUT_FILE}")
